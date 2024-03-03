import json
from ingreso.login import connect
import pandas as pd
from openpyxl import Workbook
from openpyxl.drawing.image import Image
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment, PatternFill, Border, Side
from docx import Document
from datetime import datetime
import locale

def insert_course(archivo):    
    client = connect()
    df = pd.read_excel(f'static/course-students/{archivo}')
    course_number = int(df.iloc[0, 2].split(' - ')[0])
    program = df.iloc[0, 2].split(' - ')[1]
    status = df.iloc[1, 2]
    if client:
        db = client["sara"]
        collection = db["courses"]
        course = collection.find_one({"course_number": course_number})
        if not course:
            data = {
                "course_number": course_number,
                "program": program,
                "status": status,
            }
            try:
                collection.insert_one(data)
                return True
            except:
                return False         
    return False

def insert_courses_students(archivo):
    client = connect()
    df = pd.read_excel(f'static/course-students/{archivo}')
    course_number = int(df.iloc[0, 2].split(' - ')[0])
    df = df.drop([0, 1, 2, 3])
    df.columns=['document_type', 'document', 'name', 'last_name', 'phone', 'email', 'state']
    df['course_number'] = course_number
    if client:
        db = client["sara"]
        collection = db["courses_students"]
        for index, row in df.iterrows():
            data = {
                "document_type": row['document_type'],
                "document": row['document'],
                "name": row['name'],
                "last_name": row['last_name'],
                "phone": row['phone'],
                "email": row['email'],
                "state": row['state'],
                "course_number": row['course_number']
            }
            try:
                collection.insert_one(data)
            except:
                return False
    return False


def insert_courses_competences():
    client = connect()
    with open('static/competencias.json', 'r', encoding='utf-8') as archivo_json:    
        data = json.load(archivo_json)
        df = pd.DataFrame(data)
        
        if client:
            db = client["sara"]
            collection = db["courses_competences"]            
            collection.delete_many({})            
            
            for index, row in df.iterrows():
                data = {
                    "document_type": row['tipo_documento'],
                    "document": row['numero_documento'],
                    "name": row['nombre'],
                    "last_name": row['apellidos'],
                    "status": row['estado'],
                    "competence": row['competencia'],
                    "rap": row['rap'],
                    "judgment": row['juicio'],
                    "official": row['funcionario'],
                    "course_number": row['ficha'],
                    "program": row['programa']
                }
                collection.insert_one(data)
            return True
    return False

def generate_excel_students(course_number):    
    client = connect()
    if client:
        db = client["sara"]
        collection = db["courses_students"]
        collection_courses = db["courses"]
        students = collection.find({"course_number": course_number})
        course = collection_courses.find_one({"course_number": course_number})
        students = [student for student in students if student['state'] != 'RETIRO VOLUNTARIO' and student['state'] != 'CANCELADO' and student['state'] != 'TRASLADADO']
        students = sorted(students, key=lambda x: x['name'])
        data = []
        for student in students:
            data.append(student)
        df = pd.DataFrame(data)
        df.drop(['_id', 'course_number'],  axis=1, inplace=True)
        df.columns = ['Tipo documento', 'Documento', 'Nombre', 'Apellido', 'Celular', 'Correo', 'Estado']
        df['Nombre completo'] = df['Nombre'] + ' ' + df['Apellido']
        df.drop(['Nombre', 'Apellido'], axis=1, inplace=True)
        df = df[['Tipo documento', 'Documento', 'Nombre completo', 'Celular', 'Correo', 'Estado']]
        workbook = Workbook()
        hoja = workbook.active
        img = Image('static/images/logo-sena.png')
        hoja.add_image(img, f'A1')
        hoja['B2'] = 'Centro de Servicios y Gestión Empresarial'
        hoja['B3'] = 'Coordinación de Teleinformática'
        hoja['A4'] = ''
        hoja['A5'] = 'Ficha:'
        hoja['B5'] = course_number
        hoja['A6'] = 'Programa:'
        hoja['B6'] = course['program']
        hoja['A7'] = ''
        hoja[f'A5'].font = hoja[f'A5'].font.copy(bold=True)
        hoja[f'A6'].font = hoja[f'A6'].font.copy(bold=True)    

        for i in range(2, 4):
            hoja[f'B{i}'].font = hoja[f'C{i}'].font.copy(bold=True)

        hoja['A8'] = 'Tipo documento'
        hoja['B8'] = 'Documento'
        hoja['C8'] = 'Nombre completo'
        hoja['D8'] = 'Celular'
        hoja['E8'] = 'Correo'
        hoja['F8'] = 'Estado'

        rango_celdas = hoja['A8:F8']
        color_fondo = PatternFill(start_color='C5EAE8', end_color='C5EAE8', fill_type='solid')
        for row in rango_celdas:
            for cell in row:
                cell.fill = color_fondo
        hoja.column_dimensions['A'].width = 19
        hoja.column_dimensions['B'].width = 12
        hoja.column_dimensions['C'].width = 35
        hoja.column_dimensions['D'].width = 12
        hoja.column_dimensions['E'].width = 40
        hoja.column_dimensions['F'].width = 14

        for r in dataframe_to_rows(df, index=False, header=False):
            hoja.append(r)       
        
        
        workbook.save(f'static/course-students-excel/{course_number}.xlsx')



def courses_delivery(course_number, instructor_name):
    doc = Document(f"static/course_delivery/Formato_Acta_Entrega_Ficha.docx")
    course = course_number    
    locale.setlocale(locale.LC_ALL, 'es_ES.utf8')
    date_now = datetime.now().strftime('%d de %B de %Y')
    with open('static/competencias.json', 'r', encoding='utf-8') as archivo_json:    
        data = json.load(archivo_json)
        df = pd.DataFrame(data)
        df = df[df['ficha'] == int(course)]
        program = df.iloc[0, 10]

        df = df[(df['estado'] != 'CANCELADO') & (df['estado'] != 'RETIRO VOLUNTARIO') & (df['estado'] != 'TRASLADADO')]

        df = df.groupby(['tipo_documento', 'numero_documento', 'nombre', 'apellidos', 'juicio', 'estado']).size().reset_index(name='cantidad')
        df2 = df.drop_duplicates(subset='numero_documento')
        amount_status = df2.groupby(['estado']).size().reset_index(name='cantidad')
        students = df[(df['juicio'] == 'POR EVALUAR') & (df['cantidad'] == 1)]
        untested_students = df[(df['juicio'] == 'POR EVALUAR') & (df['cantidad'] > 1)]     
        not_approved_students = df[(df['juicio'] == 'NO APROBADO')]
        unproductive_students = pd.concat([untested_students, not_approved_students]).drop_duplicates(subset='numero_documento')   
   
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "[FICHA]" in run.text:
                            run.text = run.text.replace("[FICHA]", course)
                        if "[FECHA]" in run.text:
                            run.text = run.text.replace("[FECHA]", date_now)
                        if "[PROGRAMA]" in run.text:
                            run.text = run.text.replace("[PROGRAMA]", program)     
                        if "[NOMBRE_INSTRUCTOR]" in run.text:
                            run.text = run.text.replace("[NOMBRE_INSTRUCTOR]", instructor_name)
                        if "[TABLA_CANTIDAD]" in run.text:
                            run.text = run.text.replace("[TABLA_CANTIDAD]", "")
                            table = cell.add_table(rows=1, cols=2)
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = 'Estado'
                            hdr_cells[0].paragraphs[0].runs[0].bold = True
                            hdr_cells[1].text = 'Cantidad'
                            hdr_cells[1].paragraphs[0].runs[0].bold = True
                            for index, row in amount_status.iterrows():
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(row['estado'])
                                row_cells[1].text = str(row['cantidad'])
                            if students.empty:
                                cell.add_paragraph("No hay aprendices para iniciar su etapa productiva.")
                            else:
                                cell.add_paragraph("Los siguientes aprendices se encuentran aprobados para inciar la etapa productiva:")
                                table = cell.add_table(rows=1, cols=4)
                                table.style = 'Table Grid'
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = 'Tipo documento'
                                hdr_cells[0].paragraphs[0].runs[0].bold = True
                                hdr_cells[1].text = 'Documento'
                                hdr_cells[1].paragraphs[0].runs[0].bold = True
                                hdr_cells[2].text = 'Nombre'
                                hdr_cells[2].paragraphs[0].runs[0].bold = True
                                hdr_cells[3].text = 'Estado'
                                hdr_cells[3].paragraphs[0].runs[0].bold = True
                                for index, row in students.iterrows():
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(row['tipo_documento'])
                                    row_cells[1].text = str(row['numero_documento'])
                                    row_cells[2].text = str(row['nombre']) + ' ' + str(row['apellidos'])
                                    row_cells[3].text = str(row['estado'])           
                            if unproductive_students.empty:
                                run = cell.add_paragraph().add_run("\nObservaciones")
                                run.bold = True
                                cell.add_paragraph("No hay aprendices con observaciones.")
                            else:
                                run = cell.add_paragraph().add_run("\nObservaciones")
                                run.bold = True
                                cell.add_paragraph("Los siguientes aprendices NO pueden inciar su etapa productiva debido a que tienen resultados pendientes por evaluar o No aprobados:")
                                table = cell.add_table(rows=1, cols=4)
                                table.style = 'Table Grid'
                                hdr_cells = table.rows[0].cells
                                hdr_cells[0].text = 'Tipo documento'
                                hdr_cells[0].paragraphs[0].runs[0].bold = True
                                hdr_cells[1].text = 'Documento'
                                hdr_cells[1].paragraphs[0].runs[0].bold = True
                                hdr_cells[2].text = 'Nombre'
                                hdr_cells[2].paragraphs[0].runs[0].bold = True
                                hdr_cells[3].text = 'Estado'
                                hdr_cells[3].paragraphs[0].runs[0].bold = True
                                for index, row in unproductive_students.iterrows():
                                    row_cells = table.add_row().cells
                                    row_cells[0].text = str(row['tipo_documento'])
                                    row_cells[1].text = str(row['numero_documento'])
                                    row_cells[2].text = str(row['nombre']) + ' ' + str(row['apellidos'])
                                    row_cells[3].text = str(row['estado'])
                                cell.add_paragraph("\n")
                             
                        

    doc.save(f"static/course_delivery/Acta_Entrega_Ficha_{course}.docx")